import io
import os
import re
from typing import Dict, Any, List
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def generate_markdown(title: str, content: str, metadata: Dict[str, Any] = None) -> str:
    """Generates standard Markdown output."""
    md = f"# {title}\n\n"
    if metadata:
        md += "## Metadata\n"
        for k, v in metadata.items():
            md += f"- **{k.capitalize()}**: {v}\n"
        md += "\n---\n\n"
    md += content
    return md

def generate_docx(title: str, content: str, metadata: Dict[str, Any] = None) -> io.BytesIO:
    """Generates a Microsoft Word document (.docx) byte stream."""
    doc = Document()
    doc.add_heading(title, level=0)
    
    if metadata:
        doc.add_heading("Metadata", level=2)
        for k, v in metadata.items():
            p = doc.add_paragraph()
            p.add_run(f"{k.capitalize()}: ").bold = True
            p.add_run(str(v))
        doc.add_paragraph("").paragraph_format.space_after = 12
        
    # Standard splitting of markdown content into headings and text
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # Clean markdown bold/italic formatting roughly
            clean_text = line.replace('**', '').replace('*', '')
            doc.add_paragraph(clean_text)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_pdf(title: str, content: str, metadata: Dict[str, Any] = None) -> io.BytesIO:
    """Generates a PDF document byte stream using ReportLab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#1E293B'), # Deep slate
        spaceAfter=15
    )
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading2'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#2563EB'), # Blue
        spaceBefore=15,
        spaceAfter=8
    )
    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading3'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=10,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=8
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )
    
    # Add title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 10))
    
    # Add metadata if any
    if metadata:
        meta_text = "<b>Metadata:</b><br/>" + "<br/>".join([f"<b>{k.capitalize()}:</b> {v}" for k, v in metadata.items()])
        story.append(Paragraph(meta_text, body_style))
        story.append(Spacer(1, 15))
        
    # Process text content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            story.append(Paragraph(line[4:], h2_style))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], h1_style))
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], title_style))
        elif line.startswith('- ') or line.startswith('* '):
            # Clean formatting bold
            bolded_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line[2:])
            story.append(Paragraph(f"• {bolded_line}", bullet_style))
        else:
            bolded_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(bolded_line, body_style))
            
    doc.build(story)
    buffer.seek(0)
    return buffer
